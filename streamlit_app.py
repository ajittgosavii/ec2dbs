import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ec2_sql_sizing import EC2DatabaseSizingCalculator

# Configure page
st.set_page_config(
    page_title="Enterprise AWS EC2 SQL Sizing", 
    layout="wide",
    page_icon=":bar_chart:"
)

# Add custom CSS for professional styling
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6;
    }
    .sidebar .sidebar-content {
        background: #ffffff;
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
        border-radius: 4px;
    }
    .stDownloadButton>button {
        background-color: #2196F3;
        color: white;
        font-weight: bold;
        border-radius: 4px;
    }
    .stAlert {
        border-radius: 4px;
    }
    .st-bb {
        background-color: white;
    }
</style>
""", unsafe_allow_html=True)

# App header
st.title("AWS EC2 SQL Server Sizing Calculator")
st.markdown("""
This enterprise-grade tool provides EC2 sizing recommendations for SQL Server workloads based on your on-premise infrastructure metrics.
Recommendations include development, QA, staging, and production environments.
""")

# Sidebar Inputs
with st.sidebar:
    st.header("Input Parameters")
    st.markdown("Enter your current on-premise SQL Server metrics:")
    
    inputs = {
        "on_prem_cores": st.number_input("CPU Cores", min_value=1, value=16, 
                                         help="Total number of CPU cores in your on-premise server"),
        "peak_cpu_percent": st.slider("Peak CPU Utilization (%)", 0, 100, 65, 
                                      help="Highest observed CPU utilization percentage"),
        "on_prem_ram_gb": st.number_input("RAM (GB)", min_value=1, value=64, 
                                          help="Total physical RAM in the server"),
        "peak_ram_percent": st.slider("Peak RAM Utilization (%)", 0, 100, 75, 
                                      help="Highest observed RAM utilization percentage"),
        "storage_current_gb": st.number_input("Current Storage (GB)", min_value=1, value=500, 
                                              help="Current database storage size"),
        "storage_growth_rate": st.number_input("Annual Growth Rate", min_value=0.0, max_value=1.0, value=0.15, step=0.01, 
                                              format="%.2f", help="Expected annual storage growth (e.g., 0.15 for 15%)"),
        "peak_iops": st.number_input("Peak IOPS", min_value=1, value=8000, 
                                     help="Highest observed Input/Output Operations Per Second"),
        "peak_throughput_mbps": st.number_input("Peak Throughput (MB/s)", min_value=1, value=400, 
                                                help="Highest observed data transfer rate"),
        "years": st.slider("Growth Projection (Years)", 1, 10, 3, 
                           help="Number of years to plan for future growth"),
        "workload_profile": st.selectbox("Workload Profile", 
                                        ["general", "memory", "compute"],
                                        help="""\n**Workload Type Guidelines**  \n- General: Balanced workloads like mixed OLTP and reporting  \n- Memory: Data warehouses, analytics, in-memory DBs  \n- Compute: OLTP, heavy transaction processing, CPU-bound jobs\n"""),
        "prefer_amd": st.checkbox("Include AMD Instances (Cost Optimized)", value=True,
                                 help="AMD instances are typically 10-20% cheaper than comparable Intel instances")
    }

# Initialize calculator
calculator = EC2DatabaseSizingCalculator()
calculator.inputs.update(inputs)

# Main app
if st.button("Generate Recommendations", key="generate_btn"):
    with st.spinner("Calculating EC2 sizing recommendations..."):
        try:

            results = calculator.generate_all_recommendations()

            # Remove redundant 'environment' key to avoid duplication
            for val in results.values():
                val.pop("environment", None)

            # Create DataFrame
            df = pd.DataFrame.from_dict(results, orient='index').reset_index()
            df.rename(columns={"index": "Environment"}, inplace=True)

            # Reorder columns for clarity
            preferred_order = [
                "Environment", "instance_type", "vCPUs", "RAM_GB", "storage_GB",
                "ebs_type", "iops_required", "throughput_required", "family", "processor"
            ]
            df = df[[col for col in preferred_order if col in df.columns]]

            # Display results
            st.success("✅ EC2 Sizing Recommendations Generated")

            st.markdown(
                "<style>div[data-testid='stDataFrame'] { margin: auto; width: 90%; }</style>", 
                unsafe_allow_html=True
            )
            st.dataframe(
                df.style.format({
                    "vCPUs": "{:.0f}",
                    "RAM_GB": "{:.0f}",
                    "storage_GB": "{:.0f}"
                }), use_container_width=True
            )
                df.style.format({
                    "vCPUs": "{:.0f}",
                    "RAM_GB": "{:.0f}",
                    "storage_GB": "{:.0f}"
                    }), use_container_width=True)
                }),
                use_container_width=True
            )
                "vCPUs": "{:.0f}",
                "RAM_GB": "{:.0f}",
                "storage_GB": "{:.0f}"
            }), use_container_width=True)
            
            # Export options
            st.subheader("Export Results")
            
            # CSV export
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button("Download CSV", csv, "ec2_sizing.csv", "text/csv")
            
            # DOCX export
            def create_docx_report(df):
                doc = Document()
                
                # Title
                title = doc.add_paragraph()
                title_run = title.add_run("AWS EC2 SQL Server Sizing Report")
                title_run.font.size = Pt(18)
                title_run.font.bold = True
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata
                doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Input Parameters:")
                doc.add_paragraph(f"  CPU Cores: {inputs['on_prem_cores']}")
                doc.add_paragraph(f"  Peak CPU: {inputs['peak_cpu_percent']}%")
                doc.add_paragraph(f"  RAM: {inputs['on_prem_ram_gb']}GB")
                doc.add_paragraph(f"  Peak RAM: {inputs['peak_ram_percent']}%")
                doc.add_paragraph(f"  Storage: {inputs['storage_current_gb']}GB")
                doc.add_paragraph(f"  Growth Rate: {inputs['storage_growth_rate']*100:.1f}%")
                doc.add_paragraph(f"  Projection Years: {inputs['years']}")
                
                # Table
                doc.add_heading("Recommendations", level=1)
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = col
                
                # Data
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                
                # Recommendations
                doc.add_heading("Implementation Notes", level=1)
                notes = [
                    "PROD environments should use Multi-AZ deployments for high availability",
                    "Use gp3 EBS volumes for cost-effective general storage",
                    "Use io2 EBS volumes for high-performance needs (>16K IOPS or >1GB/s throughput)",
                    "Enable EBS encryption at rest for all environments",
                    "Implement regular snapshot backups with retention policies",
                    "Monitor performance metrics with Amazon CloudWatch",
                    "Consider Reserved Instances for PROD for cost savings"
                ]
                
                for note in notes:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(note)
                
                return doc
            
            doc = create_docx_report(df)
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            st.download_button("Download DOCX Report", doc_io, 
                              "ec2_sizing_report.docx", 
                              "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # Cost optimization note
            if inputs["prefer_amd"]:
                st.info("💡 **Cost Optimization Tip**: AMD-based instances (m6a, r6a, c6a) typically offer 10-20% better price/performance than comparable Intel instances.")
            else:
                st.info("ℹ️ AMD instances excluded from recommendations as per your selection")

        except Exception as e:
            st.error(f"Error generating recommendations: {str(e)}")
            st.exception(e)

# Footer
st.markdown("---")
st.markdown("""
**Enterprise Features:**
- Environment-specific sizing (DEV, QA, SQA, PROD)
- AMD instance optimization for cost savings
- Storage growth projections
- I/O requirements calculation
- Professional reporting (CSV, DOCX)
""")